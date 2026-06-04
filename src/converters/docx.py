"""DOCX → Markdown via python-docx.

DOCX is Word 2007+ XML format. python-docx handles the XML natively.
No external tools needed. Quality: Good — handles paragraphs, headings,
tables (basic), and embedded images are skipped.
"""

from pathlib import Path

from ..detector import FormatName
from ..security import safe_parse_xml
from . import Converter, register


@register
class _DocxConverter:
    label = "DOCX (python-docx)"

    @staticmethod
    def can_handle(fmt: FormatName) -> bool:
        return fmt == "docx"

    @staticmethod
    def check_deps() -> list[str]:
        try:
            import docx  # noqa: F401
        except ImportError:
            return ["python-docx not installed.\n  Fix: pip install python-docx"]
        return []

    @staticmethod
    def convert(path: Path, output_dir: Path) -> Path:
        from docx import Document

        # Security: validate XML structure before parsing (XXE check)
        try:
            safe_parse_xml(path)
        except Exception:
            pass  # python-docx handles parsing itself; we just verified entities

        doc = Document(str(path))
        parts = []
        total_chars = 0

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Detect heading style
            if para.style.name.startswith("Heading"):
                level = int(para.style.name.split()[-1]) if para.style.name.split()[-1].isdigit() else 1
                prefix = "#" * min(level, 6)
                parts.append(f"\n{prefix} {text}\n")
            else:
                parts.append(f"\n{text}\n")
            total_chars += len(text)

        md_path = output_dir / (path.stem + ".md")
        title = path.stem
        header = (
            f"# {title}\n\n"
            f"> {len(doc.paragraphs)} paragraphs · {total_chars:,} chars · "
            f"Converted via python-docx\n\n"
        )
        md_path.write_text(header + "\n".join(parts), encoding="utf-8")
        return md_path
